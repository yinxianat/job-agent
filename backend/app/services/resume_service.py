"""
Resume file processing:
  - Extract text from PDF / DOCX / DOC
  - Write tailored resume to DOCX and PDF  (well-formatted)
  - File naming convention: JobTitle_Location_Company.ext
"""

import os
import re
import io
from pathlib import Path

import PyPDF2
import pdfplumber
import openpyxl
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor


# ── Helpers ────────────────────────────────────────────────────────────────────

def _sanitize(name: str) -> str:
    """Remove characters not safe for filenames."""
    return re.sub(r'[\\/*?:"<>|]', "", name).replace(" ", "_")


def build_filename(job_title: str, location: str, company: str) -> str:
    from datetime import date
    date_str = date.today().strftime("%Y%m%d")
    parts = [_sanitize(job_title), _sanitize(company), date_str, "resume"]
    return "_".join(p for p in parts if p)


def _is_section_header(line: str) -> bool:
    """ALL-CAPS line with at least 3 chars — used as section heading."""
    s = line.strip()
    return (
        len(s) >= 3
        and len(s) <= 60
        and s == s.upper()
        and bool(re.search(r'[A-Z]', s))
        and not re.fullmatch(r'[\d\s\W]+', s)   # skip pure numbers/punctuation
    )


def _is_bullet(line: str) -> bool:
    s = line.strip()
    return s.startswith('•') or s.startswith('-') or s.startswith('*')


# Matches "Left text    Jan 2020 – Present" or "Left text    2019 – 2023"
# The left part and date are separated by 3+ spaces (Claude's right-flush format).
_DATE_RANGE_RE = re.compile(
    r'^(.+?)\s{3,}'
    r'((?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|'
    r'Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)'
    r'\.?\s+\d{4}|\d{4})'
    r'\s*[-–—]\s*'
    r'((?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|'
    r'Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)'
    r'\.?\s+\d{4}|\d{4}|[Pp]resent|[Cc]urrent)\s*$',
    re.IGNORECASE,
)


def _parse_date_line(line: str):
    """Return (left_text, date_string) if the line contains a right-aligned date range, else None."""
    m = _DATE_RANGE_RE.match(line.strip())
    if m:
        return m.group(1).strip(), f"{m.group(2)} – {m.group(3)}"
    return None


def _add_right_tab_stop(para, position_twips: int = 9720):
    """Add a right-aligned tab stop to a DOCX paragraph (9720 twips ≈ 6.75 inches)."""
    pPr = para._p.get_or_add_pPr()
    tabs_elem = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), str(position_twips))
    tabs_elem.append(tab)
    pPr.append(tabs_elem)


def _find_header_block(lines):
    """
    Returns (name_idx, contact_idxs) — the indices of the name line and
    contact/summary lines that appear BEFORE the first ALL-CAPS section header.
    """
    first_section = len(lines)
    for i, line in enumerate(lines):
        if _is_section_header(line):
            first_section = i
            break

    name_idx = None
    contact_idxs = []
    for i in range(first_section):
        stripped = lines[i].strip()
        if not stripped:
            continue
        if name_idx is None:
            name_idx = i
        else:
            contact_idxs.append(i)

    return name_idx, set(contact_idxs)


# ── Text extraction ─────────────────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
    except Exception:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            text += (page.extract_text() or "") + "\n"
    return text.strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    return "\n".join(para.text for para in doc.paragraphs).strip()


def extract_text_from_xlsx(file_bytes: bytes) -> str:
    """Extract all cell text from an Excel workbook, row by row, sheet by sheet."""
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    lines = []
    for sheet in wb.worksheets:
        sheet_lines = []
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                sheet_lines.append("  |  ".join(cells))
        if sheet_lines:
            lines.append(f"[Sheet: {sheet.title}]")
            lines.extend(sheet_lines)
    return "\n".join(lines).strip()


def extract_resume_text(file_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    if ext in (".docx", ".doc"):
        return extract_text_from_docx(file_bytes)
    if ext in (".xlsx", ".xls"):
        return extract_text_from_xlsx(file_bytes)
    return file_bytes.decode("utf-8", errors="ignore")


def extract_job_log_text(file_bytes: bytes, filename: str) -> str:
    """Extract text from a job log file (PDF, DOCX, DOC, XLSX, XLS, TXT)."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    if ext in (".docx", ".doc"):
        return extract_text_from_docx(file_bytes)
    if ext in (".xlsx", ".xls"):
        return extract_text_from_xlsx(file_bytes)
    # Plain text, CSV, etc.
    return file_bytes.decode("utf-8", errors="ignore")


# ── PDF generation ─────────────────────────────────────────────────────────────

def _build_pdf_story(lines, name_idx, contact_idxs, compact: bool = False):
    """
    Build a ReportLab story from resume lines.
    compact=True uses tighter font sizes and spacing to fit within 2 pages.
    """
    DARK_NAVY   = HexColor('#1a2744')
    SECTION_CLR = HexColor('#2c3e50')
    MID_GRAY    = HexColor('#555555')
    TEXT_CLR    = HexColor('#222222')
    RULE_CLR    = HexColor('#3d5a80')
    LIGHT_RULE  = HexColor('#bbbbbb')

    CONTENT_WIDTH = 6.75 * inch   # 8.5" - 2×0.875" margins

    # Compact mode reduces font sizes and tightens vertical spacing
    body_fs   = 9.5  if compact else 10
    body_lead = 12   if compact else 14
    sec_fs    = 10   if compact else 10.5
    sec_sb    = 9    if compact else 14
    name_fs   = 16   if compact else 18

    name_style = ParagraphStyle(
        'ResumeName',
        fontName='Helvetica-Bold',
        fontSize=name_fs,
        leading=name_fs + 4,
        alignment=1,        # centered
        textColor=DARK_NAVY,
        spaceAfter=3,
    )
    contact_style = ParagraphStyle(
        'ResumeContact',
        fontName='Helvetica',
        fontSize=8.5 if compact else 9,
        leading=12 if compact else 13,
        alignment=1,        # centered
        textColor=MID_GRAY,
        spaceAfter=1,
    )
    section_style = ParagraphStyle(
        'ResumeSection',
        fontName='Helvetica-Bold',
        fontSize=sec_fs,
        leading=sec_fs + 3,
        textColor=SECTION_CLR,
        spaceBefore=sec_sb,
        spaceAfter=2,
    )
    bullet_style = ParagraphStyle(
        'ResumeBullet',
        fontName='Helvetica',
        fontSize=body_fs,
        leading=body_lead,
        leftIndent=14,
        firstLineIndent=0,
        spaceAfter=1,
        textColor=TEXT_CLR,
    )
    normal_style = ParagraphStyle(
        'ResumeNormal',
        fontName='Helvetica',
        fontSize=body_fs,
        leading=body_lead,
        spaceAfter=1,
        textColor=TEXT_CLR,
    )
    subtitle_style = ParagraphStyle(
        'ResumeSubtitle',
        fontName='Helvetica-BoldOblique',   # bold + italic — company / location subtitle
        fontSize=body_fs,
        leading=body_lead,
        spaceAfter=1,
        textColor=TEXT_CLR,
    )
    date_left_style = ParagraphStyle(
        'ResumeDateLeft',
        fontName='Helvetica-Bold',
        fontSize=sec_fs,
        leading=body_lead,
        textColor=TEXT_CLR,
    )
    date_right_style = ParagraphStyle(
        'ResumeDateRight',
        fontName='Helvetica-BoldOblique',   # bold + italic
        fontSize=body_fs,
        leading=body_lead,
        alignment=2,           # right-aligned
        textColor=HexColor('#222222'),
    )

    story = []
    prev_was_date_entry = False

    for i, line in enumerate(lines):
        stripped = line.strip()

        if not stripped:
            story.append(Spacer(1, 3 if compact else 4))
            # blank lines do not clear the "prev_was_date_entry" flag
            continue

        if i == name_idx:
            prev_was_date_entry = False
            story.append(Paragraph(stripped, name_style))
            story.append(HRFlowable(
                width="100%", thickness=1.5, color=RULE_CLR, spaceAfter=3, spaceBefore=0,
            ))
            continue

        if i in contact_idxs:
            prev_was_date_entry = False
            story.append(Paragraph(stripped, contact_style))
            continue

        if _is_section_header(stripped):
            prev_was_date_entry = False
            story.append(Paragraph(stripped, section_style))
            story.append(HRFlowable(
                width="100%", thickness=0.5, color=LIGHT_RULE, spaceAfter=4, spaceBefore=0,
            ))
            continue

        if _is_bullet(stripped):
            prev_was_date_entry = False
            bullet_text = stripped.lstrip('•-* ').strip()
            story.append(Paragraph(f'\u2022\u00a0{bullet_text}', bullet_style))
            continue

        # ── Job/education entry with right-aligned date ────────────────────────
        date_parts = _parse_date_line(stripped)
        if date_parts:
            prev_was_date_entry = True
            left_text, date_text = date_parts
            tbl = Table(
                [[Paragraph(left_text, date_left_style),
                  Paragraph(date_text, date_right_style)]],
                colWidths=[CONTENT_WIDTH * 0.68, CONTENT_WIDTH * 0.32],
            )
            tbl.setStyle(TableStyle([
                ('ALIGN',         (0, 0), (0, 0), 'LEFT'),
                ('ALIGN',         (1, 0), (1, 0), 'RIGHT'),
                ('VALIGN',        (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING',   (0, 0), (-1, -1), 0),
                ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
                ('TOPPADDING',    (0, 0), (-1, -1), 0),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ]))
            story.append(tbl)
            continue

        # ── Subtitle: non-bullet line immediately following a date entry ──────
        # e.g. company name or location on its own line → bold-italic
        if prev_was_date_entry:
            prev_was_date_entry = False
            story.append(Paragraph(stripped, subtitle_style))
            continue

        prev_was_date_entry = False
        story.append(Paragraph(stripped, normal_style))

    return story


def _build_cover_letter_pdf_story(lines, text_color, body_fs=10, body_lead=15):
    """Build a ReportLab story for a cover letter — plain left-aligned prose."""
    para_style = ParagraphStyle(
        'CLParagraph',
        fontName='Helvetica',
        fontSize=body_fs,
        leading=body_lead,
        textColor=text_color,
        spaceAfter=8,
        alignment=0,   # LEFT
    )
    story = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 6))
        else:
            story.append(Paragraph(stripped, para_style))
    return story


def write_pdf(tailored_text: str, output_path: str, is_cover_letter: bool = False):
    """
    Write a well-formatted resume (or cover letter) PDF using ReportLab.
    When is_cover_letter=True, renders plain left-aligned prose instead of
    the resume-specific multi-column/section layout.
    For resumes, automatically retries with compact settings if > 2 pages.
    """
    TEXT_CLR = HexColor('#222222')

    if is_cover_letter:
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            leftMargin=1.0 * inch,
            rightMargin=1.0 * inch,
            topMargin=1.0 * inch,
            bottomMargin=1.0 * inch,
        )
        lines = tailored_text.splitlines()
        story = _build_cover_letter_pdf_story(lines, TEXT_CLR)
        doc.build(story)
        return

    def _build(compact: bool):
        top_margin = (0.65 if compact else 0.75) * inch
        bot_margin = (0.65 if compact else 0.75) * inch
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            leftMargin=0.875 * inch,
            rightMargin=0.875 * inch,
            topMargin=top_margin,
            bottomMargin=bot_margin,
        )
        lines = tailored_text.splitlines()
        name_idx, contact_idxs = _find_header_block(lines)
        story = _build_pdf_story(lines, name_idx, contact_idxs, compact=compact)
        doc.build(story)

    # First pass — normal sizing
    _build(compact=False)

    # Check page count; if > 2, retry with compact settings
    try:
        reader = PyPDF2.PdfReader(output_path)
        if len(reader.pages) > 2:
            _build(compact=True)
    except Exception:
        pass  # If the check fails, keep the first build


# ── DOCX generation ────────────────────────────────────────────────────────────

def _add_paragraph_bottom_border(para):
    """Add a thin bottom border to a DOCX paragraph (section divider style)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)


def write_docx(tailored_text: str, output_path: str, is_cover_letter: bool = False):
    """
    Write a well-formatted resume or cover letter DOCX using python-docx.
    When is_cover_letter=True, renders plain left-aligned prose instead of
    the resume-specific section/header layout.
    """
    lines = tailored_text.splitlines()

    doc = Document()

    if is_cover_letter:
        # Cover letter: generous margins, plain left-aligned paragraphs
        for section in doc.sections:
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin   = Inches(1.0)
            section.right_margin  = Inches(1.0)
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)
        for line in lines:
            stripped = line.strip()
            if not stripped:
                p = doc.add_paragraph()
                p.paragraph_format.space_after  = Pt(0)
                p.paragraph_format.space_before = Pt(0)
            else:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_after = Pt(6)
                run = para.add_run(stripped)
                run.font.size = Pt(11)
        doc.save(output_path)
        return

    name_idx, contact_idxs = _find_header_block(lines)

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.875)
        section.right_margin  = Inches(0.875)

    # Default body font
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10.5)

    prev_was_date_entry = False

    for i, line in enumerate(lines):
        stripped = line.strip()

        if not stripped:
            # Minimal spacer paragraph; blank lines don't clear prev_was_date_entry
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            continue

        # ── Name ──────────────────────────────────────────────────────────────
        if i == name_idx:
            prev_was_date_entry = False
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(stripped)
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1a, 0x27, 0x44)
            # Add bottom border under name
            _add_paragraph_bottom_border(para)
            continue

        # ── Contact / header block ─────────────────────────────────────────────
        if i in contact_idxs:
            prev_was_date_entry = False
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(1)
            run = para.add_run(stripped)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # ── Section header (ALL-CAPS) ──────────────────────────────────────────
        if _is_section_header(stripped):
            prev_was_date_entry = False
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after  = Pt(3)
            run = para.add_run(stripped)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
            _add_paragraph_bottom_border(para)
            continue

        # ── Bullet ─────────────────────────────────────────────────────────────
        if _is_bullet(stripped):
            prev_was_date_entry = False
            bullet_text = stripped.lstrip('•-* ').strip()
            try:
                para = doc.add_paragraph(style='List Bullet')
            except KeyError:
                para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(1)
            para.paragraph_format.left_indent = Inches(0.2)
            run = para.add_run(bullet_text)
            run.font.size = Pt(10)
            continue

        # ── Job/education entry with right-aligned date ────────────────────────
        date_parts = _parse_date_line(stripped)
        if date_parts:
            prev_was_date_entry = True
            left_text, date_text = date_parts
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(1)
            _add_right_tab_stop(para)          # right tab at ~6.75"
            run_left = para.add_run(left_text)
            run_left.bold = True
            run_left.font.size = Pt(10.5)
            para.add_run('\t')                 # jump to right-aligned tab
            run_date = para.add_run(date_text)
            run_date.font.size = Pt(10)
            run_date.bold = True
            run_date.italic = True
            run_date.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            continue

        # ── Subtitle: non-bullet line immediately following a date entry ──────
        # e.g. company name or location on its own line → bold-italic
        if prev_was_date_entry:
            prev_was_date_entry = False
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(1)
            run = para.add_run(stripped)
            run.bold = True
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            continue

        # ── Normal line ────────────────────────────────────────────────────────
        prev_was_date_entry = False
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(1)
        run = para.add_run(stripped)
        run.font.size = Pt(10)

    doc.save(output_path)


# ── Saving files ───────────────────────────────────────────────────────────────

def save_tailored_resume(
    tailored_text: str,
    output_folder: str,
    job_title: str,
    location: str,
    company: str,
) -> tuple[str, str]:
    """Save tailored resume as DOCX + PDF. Returns (docx_path, pdf_path)."""
    os.makedirs(output_folder, exist_ok=True)
    base_name = build_filename(job_title, location, company)
    docx_path = os.path.join(output_folder, f"{base_name}.docx")
    pdf_path  = os.path.join(output_folder, f"{base_name}.pdf")
    write_docx(tailored_text, docx_path)
    write_pdf(tailored_text,  pdf_path)
    return docx_path, pdf_path


def save_cover_letter(
    cover_letter_text: str,
    output_folder: str,
    job_title: str,
    location: str,
    company: str,
) -> tuple[str, str]:
    """Save cover letter as DOCX + PDF. Returns (docx_path, pdf_path)."""
    os.makedirs(output_folder, exist_ok=True)
    base_name = "CoverLetter_" + build_filename(job_title, location, company)
    docx_path = os.path.join(output_folder, f"{base_name}.docx")
    pdf_path  = os.path.join(output_folder, f"{base_name}.pdf")
    write_docx(cover_letter_text, docx_path, is_cover_letter=True)
    write_pdf(cover_letter_text,  pdf_path,  is_cover_letter=True)
    return docx_path, pdf_path


def save_from_preview(
    tailored_text: str,
    output_folder: str,
    base_filename: str,
) -> tuple[str, str]:
    """Save an edited preview as DOCX + PDF. Returns (docx_path, pdf_path)."""
    os.makedirs(output_folder, exist_ok=True)
    docx_path = os.path.join(output_folder, f"{base_filename}.docx")
    pdf_path  = os.path.join(output_folder, f"{base_filename}.pdf")
    write_docx(tailored_text, docx_path)
    write_pdf(tailored_text,  pdf_path)
    return docx_path, pdf_path
